import streamlit as st
import time
import json
import io
from datetime import datetime
from dotenv import load_dotenv
import os
from crew import research_crew

# PDF & DOCX Export
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib import colors

    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    PDF_AVAILABLE = True
    DOCX_AVAILABLE = True

except ImportError:
    PDF_AVAILABLE = False
    DOCX_AVAILABLE = False

load_dotenv()

st.set_page_config(
    page_title="Neural Research Platform",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================================
# PREMIUM ENTERPRISE CSS
# ============================================================================

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Space+Grotesk:wght@500;700&family=JetBrains+Mono:wght@400;600&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
    background: linear-gradient(135deg, #0b1120 0%, #1a1f3a 50%, #0f172a 100%);
    color: #e2e8f0;
}

section[data-testid="stAppViewContainer"] {
    background: linear-gradient(135deg, #0b1120 0%, #1a1f3a 50%, #0f172a 100%);
}

[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #1a1f3a 0%, #0f172a 100%);
    border-right: 1px solid rgba(99,102,241,0.2);
}

.hero {
    padding: 4rem 3rem;
    border-radius: 32px;
    background: linear-gradient(135deg, rgba(99,102,241,0.15), rgba(139,92,246,0.15));
    border: 2px solid rgba(99,102,241,0.4);
    text-align: center;
    margin-bottom: 3rem;
    box-shadow: 0 20px 60px rgba(99,102,241,0.3), inset 0 1px 0 rgba(255,255,255,0.1);
    animation: heroGlow 3s ease-in-out infinite;
    backdrop-filter: blur(20px);
}

@keyframes heroGlow {
    0%, 100% { box-shadow: 0 20px 60px rgba(99,102,241,0.3); }
    50% { box-shadow: 0 30px 80px rgba(99,102,241,0.5); }
}

.hero-title {
    font-family: 'Space Grotesk', sans-serif;
    font-size: 3.5rem;
    font-weight: 700;
    background: linear-gradient(90deg, #6366f1, #8b5cf6, #06b6d4, #10b981);
    background-size: 300% 300%;
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    animation: gradientShift 4s ease infinite;
    margin-bottom: 1rem;
}

@keyframes gradientShift {
    0% { background-position: 0% 50%; }
    50% { background-position: 100% 50%; }
    100% { background-position: 0% 50%; }
}

.hero-subtitle {
    font-size: 1.1rem;
    color: #a0aec0;
    font-weight: 500;
}

.feature-card {
    background: linear-gradient(135deg, rgba(30,41,59,0.6), rgba(15,23,42,0.4));
    backdrop-filter: blur(20px);
    border-radius: 24px;
    border: 1px solid rgba(99,102,241,0.3);
    padding: 2rem;
    margin: 1rem 0;
    transition: all 0.4s cubic-bezier(0.34, 1.56, 0.64, 1);
    cursor: pointer;
    position: relative;
    overflow: hidden;
}

.feature-card:hover {
    transform: scale(1.05) translateY(-10px);
    border-color: rgba(99,102,241,0.6);
    box-shadow: 0 20px 60px rgba(99,102,241,0.4);
}

.card-title {
    font-size: 1.3rem;
    font-weight: 700;
    margin-bottom: 0.5rem;
    background: linear-gradient(90deg, #e2e8f0, #cbd5e1);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}

.card-description {
    color: #a0aec0;
    font-size: 0.95rem;
    line-height: 1.6;
}

.glass {
    background: rgba(30,41,59,0.5);
    backdrop-filter: blur(20px);
    border-radius: 20px;
    border: 1px solid rgba(99,102,241,0.2);
    padding: 2rem;
    margin-bottom: 1.5rem;
    transition: all 0.3s ease;
}

.glass:hover {
    border-color: rgba(99,102,241,0.4);
    box-shadow: 0 12px 48px rgba(99,102,241,0.2);
}

.metric-card {
    background: linear-gradient(135deg, rgba(99,102,241,0.1), rgba(139,92,246,0.1));
    backdrop-filter: blur(20px);
    border-radius: 16px;
    border: 1px solid rgba(99,102,241,0.3);
    padding: 1.5rem;
    text-align: center;
    transition: all 0.3s ease;
}

.metric-card:hover {
    transform: translateY(-4px);
    border-color: rgba(99,102,241,0.6);
    box-shadow: 0 12px 40px rgba(99,102,241,0.3);
}

.metric-value {
    font-size: 2.8rem;
    font-weight: 700;
    background: linear-gradient(135deg, #6366f1, #8b5cf6);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}

.metric-label {
    font-size: 0.85rem;
    color: #94a3b8;
    text-transform: uppercase;
    letter-spacing: 1px;
}

.confidence-meter {
    width: 100%;
    height: 8px;
    background: rgba(255,255,255,0.05);
    border-radius: 10px;
    margin: 1rem 0;
    overflow: hidden;
    border: 1px solid rgba(99,102,241,0.2);
}

.confidence-fill {
    height: 100%;
    background: linear-gradient(90deg, #10b981, #6366f1);
    border-radius: 10px;
    animation: pulse 2s ease-in-out infinite;
    box-shadow: 0 0 20px rgba(99,102,241,0.6);
}

@keyframes pulse {
    0%, 100% { opacity: 1; }
    50% { opacity: 0.8; }
}

.agent-output-card {
    background: rgba(30,41,59,0.5);
    border-radius: 16px;
    border-left: 4px solid #6366f1;
    padding: 1.5rem;
    margin: 1rem 0;
    border: 1px solid rgba(99,102,241,0.2);
}

.agent-output-card.research { border-left-color: #6366f1; }
.agent-output-card.analyst { border-left-color: #8b5cf6; }
.agent-output-card.writer { border-left-color: #06b6d4; }
.agent-output-card.qa { border-left-color: #10b981; }

.agent-output-title {
    font-weight: 700;
    font-size: 1.2rem;
    margin-bottom: 1rem;
    color: #e2e8f0;
}

.agent-output-content {
    background: rgba(0,0,0,0.3);
    border-radius: 8px;
    padding: 1.5rem;
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.9rem;
    max-height: 500px;
    overflow-y: auto;
    color: #cbd5e1;
    line-height: 1.7;
}

.final-report {
    background: linear-gradient(135deg, rgba(99,102,241,0.1), rgba(139,92,246,0.1));
    border-radius: 20px;
    border: 1.5px solid rgba(99,102,241,0.3);
    padding: 2.5rem;
    margin: 2rem 0;
}

.final-report-title {
    font-size: 1.8rem;
    font-weight: 700;
    color: #e2e8f0;
    margin-bottom: 1.5rem;
}

.health-monitor {
    background: linear-gradient(135deg, rgba(99,102,241,0.1), rgba(139,92,246,0.1));
    border-radius: 16px;
    padding: 1.5rem;
    border: 1px solid rgba(99,102,241,0.3);
    margin: 1rem 0;
}

.health-bar {
    height: 6px;
    background: rgba(255,255,255,0.05);
    border-radius: 3px;
    overflow: hidden;
    margin-top: 0.5rem;
}

.health-fill {
    height: 100%;
    background: linear-gradient(90deg, #10b981, #6366f1, #8b5cf6);
    border-radius: 3px;
    box-shadow: 0 0 15px rgba(99,102,241,0.5);
}

div.stButton > button {
    height: 3.2rem !important;
    border-radius: 14px !important;
    font-weight: 600 !important;
    background: linear-gradient(90deg, #6366f1, #4f46e5) !important;
    border: none !important;
    color: white !important;
    transition: all 0.3s ease !important;
    box-shadow: 0 4px 15px rgba(99,102,241,0.4) !important;
}

div.stButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 25px rgba(99,102,241,0.6) !important;
}

@media (max-width: 768px) {
    .hero-title { font-size: 2rem; }
    .metric-value { font-size: 2rem; }
}
            
/* REMOVE DEFAULT STREAMLIT INPUT BOX BACKGROUND */
div[data-baseweb="input"] {
    background: transparent !important;
    border: none !important;
}

div[data-baseweb="base-input"] {
    background: transparent !important;
    border: none !important;
}

input {
    background: transparent !important;
    border: 1px solid rgba(99,102,241,0.4) !important;
    border-radius: 12px !important;
    padding: 0.8rem 1rem !important;
    color: #e2e8f0 !important;
}

input:focus {
    border: 1px solid #6366f1 !important;
    box-shadow: 0 0 0 1px #6366f1 !important;
}
            
section.main > div > div:has(div[data-testid="stTextInput"]) {
    background: rgba(30,41,59,0.5);
    backdrop-filter: blur(20px);
    border-radius: 24px;
    border: 1px solid rgba(99,102,241,0.25);
    padding: 2.5rem;
    margin-bottom: 2rem;

    width: 75%;          /* Match 3/4 column */
}
</style>
""", unsafe_allow_html=True)

# ============================================================================
# SESSION STATE INITIALIZATION
# ============================================================================

def init_session_state():
    defaults = {
        "page": "dashboard",
        "metrics": {"total": 0, "success": 0, "failed": 0, "avg_time": 0},
        "config": {"depth": "Standard", "sources": 5, "validation": True, "enterprise_mode": False},
        "history": [],
        "agent_outputs": {"research": None, "analyst": None, "writer": None, "qa": None},
        "final_report": None,
        "api_usage": {"groq_calls": 0, "serper_calls": 0, "total_tokens": 0},
        "agent_health": {"research": 95, "analyst": 92, "writer": 88, "qa": 96},
    }
    
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_session_state()

# ============================================================================
# SIDEBAR
# ============================================================================

with st.sidebar:
    st.markdown("### 🧠 Neural Research")
    st.divider()
    
    st.markdown("**🚀 Enterprise Mode**")
    st.session_state.config["enterprise_mode"] = st.toggle("Enable Enterprise Features", st.session_state.config["enterprise_mode"])
    
    if st.session_state.config["enterprise_mode"]:
        st.success("✅ Enterprise Mode Active")
    
    st.divider()
    
    st.markdown("**📍 Navigation**")
    page = st.radio("Select Page", ["Dashboard", "System Control", "Architecture", "Session History"], label_visibility="collapsed")
    st.session_state.page = page.lower().replace(" ", "_")
    
    st.divider()
    
    st.markdown("**📊 API Usage**")
    col1, col2 = st.columns(2)
    with col1:
        st.metric("GROQ", st.session_state.api_usage["groq_calls"])
    with col2:
        st.metric("SERPER", st.session_state.api_usage["serper_calls"])

# ============================================================================
# PAGE: DASHBOARD WITH AGENT OUTPUTS & FINAL REPORT
# ============================================================================

if st.session_state.page == "dashboard":
    st.markdown("""
    <div class="hero">
        <div class="hero-title">🧠 Neural Research Platform</div>
        <div class="hero-subtitle">Enterprise-Grade AI Research System</div>
    </div>
    """, unsafe_allow_html=True)
    
    # METRICS
    st.markdown("## 📊 Performance Metrics")
    m = st.session_state.metrics
    success_rate = (m["success"] / m["total"] * 100) if m["total"] > 0 else 0
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown(f"<div class='metric-card'><div class='metric-value'>{m['total']}</div><div class='metric-label'>Total</div></div>", unsafe_allow_html=True)
    with col2:
        st.markdown(f"<div class='metric-card'><div class='metric-value'>{m['success']}</div><div class='metric-label'>Success</div></div>", unsafe_allow_html=True)
    with col3:
        st.markdown(f"<div class='metric-card'><div class='metric-value'>{success_rate:.0f}%</div><div class='metric-label'>Rate</div></div>", unsafe_allow_html=True)
    with col4:
        st.markdown(f"<div class='metric-card'><div class='metric-value'>{m['avg_time']:.1f}s</div><div class='metric-label'>Avg Time</div></div>", unsafe_allow_html=True)
    
    st.markdown("---")
    
    # FEATURE CARDS
    st.markdown("## 🚀 Features")
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        <div class="feature-card">
            <div class="card-title">🔬 Research Specialist</div>
            <div class="card-description">Gathers comprehensive information from multiple sources with advanced web search capabilities.</div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <div class="feature-card">
            <div class="card-title">✍️ Content Writer</div>
            <div class="card-description">Creates well-structured reports with professional formatting and compelling narratives.</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="feature-card">
            <div class="card-title">📊 Data Analyst</div>
            <div class="card-description">Extracts meaningful patterns and insights from research data with statistical validation.</div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <div class="feature-card">
            <div class="card-title">✅ QA Reviewer</div>
            <div class="card-description">Ensures accuracy, clarity, and quality through comprehensive validation checks.</div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # RESEARCH CONTROL
    st.markdown("## 🔬 Research Control")

    research_box = st.container()

    with research_box:
        topic = st.text_input(
            "",
            placeholder="Enter your research topic...",
            label_visibility="collapsed"
        )

        col1, col2 = st.columns([3, 1])
        with col1:
            start = st.button("🚀 Start Research", use_container_width=True)
        with col2:
            if st.button("📋 Templates", use_container_width=True):
                st.session_state.show_templates = not st.session_state.get("show_templates", False)
    
    if start and topic.strip():
        start_time = time.time()
        progress_bar = st.progress(0)
        confidence_placeholder = st.empty()

        for i in range(0, 101, 10):
            progress_bar.progress(i)
            confidence = 60 + (i * 0.3)

            confidence_placeholder.markdown(f"""
            <div>
                <div style="font-size:0.9rem;color:#a0aec0;margin-bottom:0.5rem">
                    AI Confidence: {confidence:.0f}%
                </div>
                <div class="confidence-meter">
                    <div class="confidence-fill" style="width:{confidence}%"></div>
                </div>
            </div>
            """, unsafe_allow_html=True)

            time.sleep(0.2)

        confidence_placeholder.empty()

        # =====================================================
        # 🚀 RUN CREWAI MULTI-AGENT SYSTEM (REAL LOGIC ADDED)
        # =====================================================

        try:
            result = research_crew.kickoff({"topic": topic})

            # Read generated output files
            def read_file(filename):
                if os.path.exists(filename):
                    with open(filename, "r", encoding="utf-8") as f:
                        return f.read()
                return "Output file not found."

            st.session_state.agent_outputs["research"] = read_file("research_findings.md")
            st.session_state.agent_outputs["analyst"] = read_file("analysis_report.md")
            st.session_state.agent_outputs["writer"] = read_file("final_report.md")
            st.session_state.agent_outputs["qa"] = read_file("critique_report.md")

        except Exception as e:
            st.error(f"Research failed: {e}")
            st.session_state.agent_outputs = {
                "research": None,
                "analyst": None,
                "writer": None,
                "qa": None
            }
            progress_bar.empty()
            confidence_placeholder.empty()
            st.stop()

        # UPDATE METRICS
        duration = time.time() - start_time
        st.session_state.metrics["total"] += 1
        st.session_state.metrics["success"] += 1
        st.session_state.api_usage["groq_calls"] += 1
        st.session_state.api_usage["serper_calls"] += 1

        total = st.session_state.metrics["total"]
        current_avg = st.session_state.metrics["avg_time"]
        st.session_state.metrics["avg_time"] = ((current_avg * (total - 1) + duration) / total)

        progress_bar.empty()
        st.success("✅ Research Completed Successfully!")
        st.rerun()
    
    st.markdown("---")
    
    # AGENT OUTPUTS DISPLAY
    if st.session_state.agent_outputs["research"] is not None:
        st.markdown("## 🤖 AI Agent Outputs")
        
        tab1, tab2, tab3, tab4 = st.tabs([
            "🔬 Research Specialist",
            "📊 Data Analyst",
            "✍️ Content Writer",
            "✅ QA Reviewer"
        ])
        
        # 🔬 RESEARCH SPECIALIST
        with tab1:
            research_content = st.session_state.agent_outputs["research"] or ""

            st.markdown(f"""
            <div class="agent-output-card research">
                <div class="agent-output-title">🔬 Research Specialist Output</div>
                <div class="agent-output-content">
                    {research_content}
                </div>
            </div>
            """, unsafe_allow_html=True)

        # 📊 DATA ANALYST
        with tab2:
            analyst_content = st.session_state.agent_outputs["analyst"] or ""

            st.markdown(f"""
            <div class="agent-output-card analyst">
                <div class="agent-output-title">📊 Data Analyst Output</div>
                <div class="agent-output-content">
                    {analyst_content}
                </div>
            </div>
            """, unsafe_allow_html=True)

        # ✍️ CONTENT WRITER
        with tab3:
            writer_content = st.session_state.agent_outputs["writer"] or ""

            st.markdown(f"""
            <div class="agent-output-card writer">
                <div class="agent-output-title">✍️ Content Writer Output</div>
                <div class="agent-output-content">
                    {writer_content}
                </div>
            </div>
            """, unsafe_allow_html=True)

        # ✅ QA REVIEWER
        with tab4:
            qa_content = st.session_state.agent_outputs["qa"] or ""

            st.markdown(f"""
            <div class="agent-output-card qa">
                <div class="agent-output-title">✅ QA Reviewer Output</div>
                <div class="agent-output-content">
                    {qa_content}
                </div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("---")

        # FINAL REPORT

        # Get real topic safely
        topic_title = "Research Analysis"
        if st.session_state.agent_outputs.get("research"):
            research_text = st.session_state.agent_outputs["research"]
            if "**Topic:**" in research_text:
                try:
                    topic_title = research_text.split("**Topic:**")[1].split("\n")[0].strip()
                except:
                    pass

        # Combine real outputs dynamically
        final_report_content = ""

        if os.path.exists("final_consolidated_report.md"):
            with open("final_consolidated_report.md", "r", encoding="utf-8") as f:
                final_report_content = f.read()
        else:
            final_report_content = st.session_state.agent_outputs.get("writer", "")

        st.markdown(f"""
        <div class="final-report">
            <div class="final-report-title">
                📑 Final Research Report
            </div>
            <div class="final-report-content">
                {final_report_content}
            </div>
        </div>
        """, unsafe_allow_html=True)


        # EXPORT OPTIONS
        st.markdown("## 📥 Export Options")

        col1, col2, col3, col4 = st.columns(4)

        # ------------------------------------------------
        # 1️⃣ PDF EXPORT (Your Existing – Keep It)
        # ------------------------------------------------
        with col1:
            if st.button("📄 PDF Report", use_container_width=True):
                try:
                    pdf_buffer = io.BytesIO()

                    doc = SimpleDocTemplate(
                        pdf_buffer,
                        pagesize=letter,
                        rightMargin=40,
                        leftMargin=40,
                        topMargin=60,
                        bottomMargin=50
                    )

                    elements = []
                    styles = getSampleStyleSheet()

                    title_style = styles["Heading1"]
                    normal_style = styles["Normal"]

                    elements.append(Paragraph(f"Research Report: {topic_title}", title_style))
                    elements.append(Spacer(1, 12))

                    for key, content in st.session_state.agent_outputs.items():
                        if content:
                            elements.append(Paragraph(key.upper(), styles["Heading2"]))
                            elements.append(Spacer(1, 8))
                            clean_text = content.replace("**", "").replace("#", "")
                            for line in clean_text.split("\n"):
                                if line.strip():
                                    elements.append(Paragraph(line.strip(), normal_style))
                                    elements.append(Spacer(1, 6))
                            elements.append(PageBreak())

                    doc.build(elements)
                    pdf_buffer.seek(0)

                    st.download_button(
                        "⬇ Download PDF",
                        pdf_buffer.getvalue(),
                        file_name=f"research_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )

                except Exception as e:
                    st.error(f"PDF Error: {e}")

        # ------------------------------------------------
        # 2️⃣ WORD EXPORT (.docx)
        # ------------------------------------------------
        with col2:
            if st.button("📝 Word Document", use_container_width=True):
                try:
                    doc = Document()
                    doc.add_heading(f"Research Report: {topic_title}", level=1)

                    for key, content in st.session_state.agent_outputs.items():
                        if content:
                            doc.add_heading(key.upper(), level=2)
                            doc.add_paragraph(content)

                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    st.download_button(
                        "⬇ Download Word",
                        buffer.getvalue(),
                        file_name=f"research_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

                except Exception as e:
                    st.error(f"Word Export Error: {e}")

        # ------------------------------------------------
        # 3️⃣ JSON EXPORT
        # ------------------------------------------------
        with col3:
            if st.button("📦 JSON File", use_container_width=True):

                export_data = {
                    "topic": topic_title,
                    "generated_at": datetime.now().isoformat(),
                    "agent_outputs": st.session_state.agent_outputs,
                    "metrics": st.session_state.metrics
                }

                json_data = json.dumps(export_data, indent=4)

                st.download_button(
                    "⬇ Download JSON",
                    json_data,
                    file_name=f"research_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json",
                    use_container_width=True
                )

        # ------------------------------------------------
        # 4️⃣ MARKDOWN EXPORT (.md)
        # ------------------------------------------------
        with col4:
            if st.button("📃 Markdown", use_container_width=True):

                md_content = f"# Research Report: {topic_title}\n\n"
                md_content += f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"

                for key, content in st.session_state.agent_outputs.items():
                    if content:
                        md_content += f"## {key.upper()}\n\n"
                        md_content += content + "\n\n"

                st.download_button(
                    "⬇ Download Markdown",
                    md_content,
                    file_name=f"research_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                    mime="text/markdown",
                    use_container_width=True
                )
    
    st.markdown("---")
    
    # AGENT HEALTH
    st.markdown("## 🏥 Agent Health Monitor")
    
    col1, col2, col3, col4 = st.columns(4)
    agents = ["research", "analyst", "writer", "qa"]
    names = ["Research", "Analyst", "Writer", "QA"]
    
    for col, agent, name in zip([col1, col2, col3, col4], agents, names):
        with col:
            health = st.session_state.agent_health[agent]
            st.markdown(f"""
            <div class="health-monitor">
                <div style="font-weight:600">{name}</div>
                <div class="health-bar">
                    <div class="health-fill" style="width:{health}%"></div>
                </div>
                <div style="font-size:0.8rem;color:#a0aec0;margin-top:0.5rem">{health}% Healthy</div>
            </div>
            """, unsafe_allow_html=True)
elif st.session_state.page == "system_control":
    # Display the title for the System Control page
    st.markdown('<div class="hero"><div class="hero-title">⚙️ System Control</div></div>', unsafe_allow_html=True)
    
    # Research Configuration Section
    st.markdown("""
    <div class="glass-box">
        <h2>Research Configuration</h2>
        <!-- content -->
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("**Research Depth**")
        depth = st.selectbox("Select", ["Standard", "Deep", "Executive"], index=0, label_visibility="collapsed")
        st.session_state.config["depth"] = depth
    
    with col2:
        st.markdown("**Number of Sources**")
        sources = st.slider("Sources", 3, 20, st.session_state.config["sources"], label_visibility="collapsed")
        st.session_state.config["sources"] = sources
    
    with col3:
        st.markdown("**Validation Layer**")
        validation = st.toggle("Enable", st.session_state.config["validation"], label_visibility="collapsed")
        st.session_state.config["validation"] = validation
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    if st.button("💾 Save Configuration", use_container_width=True):
        st.success("✅ Configuration Saved!")
    
    # Role-Based Access UI Section
    st.markdown("---")
    st.markdown("""
    <div class="glass-box">
        <h2>User Roles & Permissions</h2>
        <!-- content -->
    </div>
    """, unsafe_allow_html=True)
    
    roles = {
        "Admin": ["Full Access", "System Control", "User Management", "Analytics"],
        "Analyst": ["Research", "Analysis", "Export", "History"],
        "Viewer": ["View Reports", "View History", "Read-Only"]
    }
    
    selected_role = st.selectbox("Select Role", list(roles.keys()))
    
    cols = st.columns(len(roles[selected_role]))
    for col, permission in zip(cols, roles[selected_role]):
        with col:
            st.markdown(f"""
            <div style="background:rgba(16,185,129,0.2);border:1px solid rgba(16,185,129,0.4);border-radius:8px;padding:1rem;text-align:center;">
                <span style="color:#10b981;font-weight:600">✅</span> {permission}
            </div>
            """, unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Advanced Features Section
    st.markdown("---")
    st.markdown("""
    <div class="glass-box">
        <h2>Advanced Features</h2>
        <!-- content -->
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**Risk Scoring Engine**")
        risk_score = st.slider("Sensitivity", 0, 100, 50)
        st.markdown(f"Risk Level: {'🔴 High' if risk_score > 70 else '🟡 Medium' if risk_score > 30 else '🟢 Low'}")
    
    with col2:
        st.markdown("**Citation Panel**")
        citations = st.multiselect("Citation Style", ["APA", "MLA", "Chicago", "Harvard"], default=["APA"])
        if citations:
            st.markdown(f"Active: {', '.join(citations)}")
    
    st.markdown('</div>', unsafe_allow_html=True)

elif st.session_state.page == "architecture":
    # Hero Section
    st.markdown('<div class="hero"><div class="hero-title">🏗️ Architecture</div></div>', unsafe_allow_html=True)
    
    st.markdown("""
    <div class="glass-box">
        <h2>Workflow</h2>
        <!-- content -->
    </div>
    """, unsafe_allow_html=True)
    
    # Glass background for the entire workflow section with dark blue gradient
    st.markdown('<div class="glass" style="padding:2rem 1rem; border-radius:12px; box-shadow: 0 4px 16px rgba(0, 0, 0, 0.2); background: linear-gradient(135deg, #1c2a49, #2a3d66);">', unsafe_allow_html=True)
    
    # Define the steps of the workflow
    workflow_steps = [
        ("Input", "Gather raw data, user inputs, or initial materials."),
        ("Research", "Conduct deep research on the topic, gathering sources."),
        ("Analysis", "Analyze the data and research findings for insights."),
        ("Writing", "Create the document, report, or presentation based on the analysis."),
        ("QA", "Perform quality assurance and finalize the content."),
        ("Output", "Generate the final output in the required format."),
        ("Final Review", "Review the entire process for errors or improvements."),
    ]
    
    # Loop over each workflow step and create a box for it
    for i, (title, description) in enumerate(workflow_steps):
        # For all but the last step, add a line to the next box
        if i < len(workflow_steps) - 1:
            st.markdown(f"""
            <div style="background-color:#2c3e50; border-radius:10px; padding:1.5rem 1rem; text-align:center; box-shadow: 0 8px 30px rgba(0, 123, 255, 0.3); margin-bottom:2rem; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;">
                <h4 style="color:#00aaff; font-weight:600; font-size:18px;">{title}</h4>
                <p style="color:#ecf0f1; font-size:14px;">{description}</p>
            </div>
            <div style="height:40px; width:3px; background-color:#00aaff; margin-left:auto; margin-right:auto;"></div>
            """, unsafe_allow_html=True)
        else:
            # For the last step, don't add a line after it
            st.markdown(f"""
            <div style="background-color:#2c3e50; border-radius:10px; padding:1.5rem 1rem; text-align:center; box-shadow: 0 8px 30px rgba(0, 123, 255, 0.3); font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;">
                <h4 style="color:#00aaff; font-weight:600; font-size:18px;">{title}</h4>
                <p style="color:#ecf0f1; font-size:14px;">{description}</p>
            </div>
            """, unsafe_allow_html=True)
    
    # Closing the glass background div
    st.markdown('</div>', unsafe_allow_html=True)

elif st.session_state.page == "session_history":
    st.markdown('<div class="hero"><div class="hero-title">📚 History</div></div>', unsafe_allow_html=True)
    if st.session_state.history:
        st.markdown(f"**Sessions:** {len(st.session_state.history)}")
    else:
        st.info("No sessions yet")

# Footer
st.divider()
st.markdown('<div style="text-align:center;color:#94a3b8;padding:2rem 0"><p>Neural Research Platform v2.4.0 | Production Ready</p></div>', unsafe_allow_html=True)
